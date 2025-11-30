from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from io import BytesIO
from PIL import Image
import typing
import traceback
import os
import json
import uuid
import numpy as np
from dotenv import load_dotenv
load_dotenv()  # Load environment variables from a .env file
API_KEY = os.getenv("GEMINI_API_KEY")

# Embeddings and vector store
try:
    from sentence_transformers import SentenceTransformer
    S2_AVAILABLE = True
except Exception:
    SentenceTransformer = None
    S2_AVAILABLE = False

try:
    import faiss
    FAISS_AVAILABLE = True
except Exception:
    faiss = None
    FAISS_AVAILABLE = False

import pandas as pd
from docx import Document

app = FastAPI(title="Handwritten OCR + RAG API")

# Allow all origins for UI development; lock this down in production
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve frontend static files (mounted at /static) and index at /
app.mount("/static", StaticFiles(directory="frontend"), name="static")


@app.get("/", response_class=FileResponse)
async def root_index():
    return FileResponse("frontend/index.html")


# SESSION STORAGE: server-side sessions directory
SESSIONS_DIR = os.path.join(os.path.dirname(__file__), 'sessions')
os.makedirs(SESSIONS_DIR, exist_ok=True)



def chunk_text(text: str, max_chars: int = 1000, overlap: int = 200):
    if not text:
        return []
    chunks = []
    start = 0
    L = len(text)
    while start < L:
        end = min(L, start + max_chars)
        chunk = text[start:end]
        chunks.append(chunk.strip())
        if end == L:
            break
        start = end - overlap
    return chunks

def embed_chunks(model, chunks: typing.List[str]):
    if not chunks:
        return np.zeros((0, model.get_sentence_embedding_dimension()), dtype='float32')
    emb = model.encode(chunks, convert_to_numpy=True, show_progress_bar=False)
    # normalize for cosine similarity with inner product
    norms = np.linalg.norm(emb, axis=1, keepdims=True)
    norms[norms == 0] = 1
    emb = emb / norms
    return emb.astype('float32')

def build_faiss_index(embeddings: np.ndarray):
    d = embeddings.shape[1]
    index = faiss.IndexFlatIP(d)
    index.add(embeddings)
    return index

def save_session_data(session_id: str, metadata: dict, chunks: typing.List[str], embeddings: np.ndarray, index):
    base = os.path.join(SESSIONS_DIR, session_id)
    os.makedirs(base, exist_ok=True)
    # save metadata
    with open(os.path.join(base, 'meta.json'), 'w', encoding='utf8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    # save chunks
    with open(os.path.join(base, 'chunks.json'), 'w', encoding='utf8') as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)
    # save embeddings
    np.save(os.path.join(base, 'embeddings.npy'), embeddings)
    # save faiss index
    faiss.write_index(index, os.path.join(base, 'index.faiss'))

def load_session_index(session_id: str):
    base = os.path.join(SESSIONS_DIR, session_id)
    if not os.path.isdir(base):
        return None
    index_path = os.path.join(base, 'index.faiss')
    if not os.path.exists(index_path):
        return None
    idx = faiss.read_index(index_path)
    with open(os.path.join(base, 'chunks.json'), 'r', encoding='utf8') as f:
        chunks = json.load(f)
    with open(os.path.join(base, 'meta.json'), 'r', encoding='utf8') as f:
        meta = json.load(f)
    return idx, chunks, meta

# Optional dependencies and fallbacks
try:
    import ollama
    OLLAMA_AVAILABLE = True
except Exception:
    ollama = None
    OLLAMA_AVAILABLE = False
try:
    import google.genai as genai
    GENAI_AVAILABLE = True
except Exception:
    genai = None
    GENAI_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    REPORTLAB_AVAILABLE = True
except Exception:
    canvas = None
    letter = (612, 792)
    REPORTLAB_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except Exception:
    convert_from_bytes = None
    PDF2IMAGE_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except Exception:
    fitz = None
    PYMUPDF_AVAILABLE = False


# Utility functions copied/adapted from test2.py
def truncate_for_context(text: str, max_chars: int = 6000) -> str:
    if len(text) <= max_chars:
        return text
    head = text[: max_chars // 2]
    tail = text[- (max_chars // 2) :]
    return head + "\n\n...TRUNCATED...\n\n" + tail


def text_to_pdf_bytes(text: str) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab not available; install with: python -m pip install reportlab")
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    textobject = c.beginText()
    textobject.setTextOrigin(40, height - 40)
    for line in text.split("\n"):
        textobject.textLine(line)
    c.drawText(textobject)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def pdf_to_images(pdf_bytes: bytes) -> typing.List[Image.Image]:
    if PDF2IMAGE_AVAILABLE:
        try:
            images = convert_from_bytes(pdf_bytes)
            return images
        except Exception:
            pass

    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            images = []
            for page in doc:
                pix = page.get_pixmap()
                mode = "RGBA" if pix.alpha else "RGB"
                img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                images.append(img)
            return images
        except Exception as e:
            raise RuntimeError(f"Failed to render PDF pages: {e}")

    raise RuntimeError(
        "No PDF rendering backend available. Install `pdf2image` + poppler or `PyMuPDF` (pip install pymupdf)."
    )


def ocr_image_bytes(img_bytes: bytes) -> str:
    # Accept image bytes (PIL-ready), convert to PNG bytes for Ollama input
    img = Image.open(BytesIO(img_bytes))
    buf = BytesIO()
    img.save(buf, format="PNG")
    png_bytes = buf.getvalue()

    if not OLLAMA_AVAILABLE:
        raise RuntimeError("ollama client not available. Install and configure Ollama or mock this in development.")

    # Call Ollama chat API similarly to test2.py
    response = ollama.chat(
        model="gemma3:12b",
        messages=[
            {"role": "user", "content": "Extract ALL text from this handwriting. Return text extracted from the iamges.", "images": [png_bytes]}
        ],
    )
    # Best-effort extraction of content
    return response.get("message", {}).get("content")

def gemini_image_ocr(img_bytes: bytes) -> str:
    if not GENAI_AVAILABLE:
        raise RuntimeError("google genai client not available. Install and configure google-genai or mock this in development.")
    
    client = genai.Client(api_key=API_KEY)
    # Convert image bytes to base64 or suitable format if needed
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents="Extract ALL text from this handwriting. Return only text.",
        images=[img_bytes]
    )
    return response.text

# Request/response models
class RAGRequest(BaseModel):
    text: str
    question: str
    use_full: bool = False


@app.get("/health")
async def health():
    return {"status": "ok", "ollama": OLLAMA_AVAILABLE, "reportlab": REPORTLAB_AVAILABLE,"genai": GENAI_AVAILABLE}


@app.post("/ocr")
async def ocr_endpoint(file: UploadFile = File(...)):
    try:
        data = await file.read()
        is_pdf = file.content_type == "application/pdf" or file.filename.lower().endswith('.pdf')
        images = []
        if is_pdf:
            try:
                images = pdf_to_images(data)
            except Exception as e:
                print(e)
        else:
            # single image
            images = [Image.open(BytesIO(data))]

        texts = []
        for i, img in enumerate(images, start=1):
            # convert PIL image to bytes and call ocr_image_bytes
            buf = BytesIO()
            img.save(buf, format="PNG")
            page_bytes = buf.getvalue()
            try:
                if GENAI_AVAILABLE:
                    page_text = gemini_image_ocr(page_bytes)
                elif OLLAMA_AVAILABLE:
                    page_text = ocr_image_bytes(page_bytes)
            except Exception as e:
                # include page-specific error but continue
                page_text = f""  # leave blank on fail
            texts.append(page_text)

        full_text = "\n\n".join(texts)
        return JSONResponse({"text": full_text, "pages": len(images)})
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + "\n" + tb)


@app.post("/process")
async def process_endpoint(file: UploadFile = File(...)):
    """Process upload: OCR -> structured extraction (LLM) -> chunking -> embeddings -> FAISS index. Returns session_id and extracted fields."""
    if not (S2_AVAILABLE and FAISS_AVAILABLE):
        raise HTTPException(status_code=500, detail="Embeddings/FAISS dependencies missing on server. Check requirements.")
    try:
        data = await file.read()
        is_pdf = file.content_type == "application/pdf" or file.filename.lower().endswith('.pdf')
        images = []
        if is_pdf:
            images = pdf_to_images(data)
        else:
            images = [Image.open(BytesIO(data))]

        # OCR pages
        texts = []
        for i, img in enumerate(images, start=1):
            page_buf = BytesIO()
            img.save(page_buf, format='PNG')
            page_bytes = page_buf.getvalue()
            page_text = ''

            if not page_text and GENAI_AVAILABLE:
                try:
                    page_text = gemini_image_ocr(page_bytes)
                except Exception:
                    page_text = ''

            if not page_text and OLLAMA_AVAILABLE:
                try:
                    page_text = ocr_image_bytes(page_bytes)
                except Exception:
                    page_text = ''
            texts.append(page_text)

        full_text = '\n\n'.join(texts)

        # Ask LLM to extract structured fields (JSON). Keep prompt strict to return JSON.
        structured = {}
        if GENAI_AVAILABLE:
            try:
                prompt = (
                    "Extract structured metadata from the following handwritten document. and also extract the full text from the image. "
                    "Return ONLY a JSON object with keys: title, author, date, institution, keywords (array), and the full text from the image. "
                    "If a field is not present, return empty string or empty array.\n\nDocument text:\n")
                
                client = genai.Client(api_key=API_KEY)
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=prompt + full_text
                )
                content = response.text
                # try to parse JSON from content
                try:
                    # sometimes LLMs wrap JSON in backticks or text; extract first brace
                    start = content.find('{')
                    end = content.rfind('}')
                    if start != -1 and end != -1:
                        jtxt = content[start:end+1]
                        structured = json.loads(jtxt)
                except Exception:
                    structured = { 'title':'', 'author':'', 'date':'', 'institution':'', 'keywords':[], 'full_text':'' }
            except Exception:
                structured = { 'title':'', 'author':'', 'date':'', 'institution':'', 'keywords':[], 'full_text':'' }



        if OLLAMA_AVAILABLE and genai is None:
            try:
                prompt = (
                    "Extract structured metadata from the following handwritten document. and also extract the full text from the image. "
                    "Return ONLY a JSON object with keys: title, author, date, institution, keywords (array), and the full text from the image. "
                    "If a field is not present, return empty string or empty array.\n\nDocument text:\n")
                
                resp = ollama.chat(model='gemma3:12b', messages=[{'role':'user','content':prompt}])
                content = resp.get('message', {}).get('content', '')
                # try to parse JSON from content
                try:
                    # sometimes LLMs wrap JSON in backticks or text; extract first brace
                    start = content.find('{')
                    end = content.rfind('}')
                    if start != -1 and end != -1:
                        jtxt = content[start:end+1]
                        structured = json.loads(jtxt)
                except Exception:
                    structured = { 'title':'', 'author':'', 'date':'', 'institution':'', 'keywords':[], 'full_text':'' }
            except Exception:
                structured = { 'title':'', 'author':'', 'date':'', 'institution':'', 'keywords':[], 'full_text':'' }
        # always include raw extracted text as full_text for easier display and RAG fallback
        try:
            structured['full_text'] = full_text
        except Exception:
            structured['full_text'] = full_text
        # chunk and embed
        chunks = chunk_text(full_text, max_chars=1000, overlap=200)
        model = SentenceTransformer('all-MiniLM-L6-v2')
        embeddings = embed_chunks(model, chunks)
        index = build_faiss_index(embeddings)

        # save session
        session_id = uuid.uuid4().hex
        metadata = {
            'filename': file.filename,
            'pages': len(images),
            'structured': structured,
            'created_at': int(__import__('time').time()),
        }
        save_session_data(session_id, metadata, chunks, embeddings, index)

        return JSONResponse({'session_id': session_id, 'meta': metadata})
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + "\n" + tb)


@app.post("/rag")
async def rag_endpoint(req: RAGRequest):
    if not GENAI_AVAILABLE:
        raise HTTPException(status_code=500, detail="google genai client not available on server")
    if not OLLAMA_AVAILABLE:
        raise HTTPException(status_code=500, detail="ollama client not available on server")
    if not req.text:
        raise HTTPException(status_code=400, detail="text is required for RAG")
    ctx = req.text if req.use_full else truncate_for_context(req.text, max_chars=5000)
    system_msg = (
        "You are an assistant answering questions using the user's extracted handwritten notes as context. "
        "Use only the provided notes to answer and cite information if it's referenced to outside the notes.\n\nNotes:\n"
        + ctx
    )
    try:
        if GENAI_AVAILABLE:
            client = genai.Client(api_key=API_KEY)
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=system_msg + "\n\nQuestion:\n" + req.question
            )
            answer = response.text
            return JSONResponse({"answer": answer})
        elif OLLAMA_AVAILABLE: 
            response = ollama.chat(
            model="gemma3:12b",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": req.question},
            ],
        )
            answer = response.get("message", {}).get("content")
            return JSONResponse({"answer": answer})
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + "\n" + tb)


class QueryRequest(BaseModel):
    session_id: str
    question: str
    top_k: int = 5
    use_full: bool = False


@app.post('/query')
async def query_endpoint(req: QueryRequest):
    if not FAISS_AVAILABLE or not S2_AVAILABLE:
        raise HTTPException(status_code=500, detail='Embeddings/FAISS not available')
    
    if not OLLAMA_AVAILABLE:
        raise HTTPException(status_code=500, detail='ollama client not configured on server')
    
    try:
        idx_chunks = load_session_index(req.session_id)
        if idx_chunks is None:
            raise HTTPException(status_code=404, detail='Session not found or missing index')
        
        idx, chunks, meta = idx_chunks
        retrieved = []  # Always define retrieved upfront
        
        if req.use_full:
            # Use full session text
            structured = meta.get('structured', {})
            full_text = ''
            if isinstance(structured, dict):
                full_text = structured.get('full_text', '')
            if not full_text:
                full_text = '\n\n'.join(chunks)
            context = full_text
            retrieved = chunks  # Consider all chunks as retrieved
        else:
            # Retrieve top-k relevant chunks via FAISS
            model = SentenceTransformer('all-MiniLM-L6-v2')
            q_emb = model.encode([req.question], convert_to_numpy=True)
            q_emb = q_emb / np.linalg.norm(q_emb, axis=1, keepdims=True)
            q_emb = q_emb.astype('float32')
            D, I = idx.search(q_emb, req.top_k)
            
            for ix in I[0]:
                if ix < len(chunks):
                    retrieved.append(chunks[ix])
            context = '\n\n'.join(retrieved) if retrieved else ''
        
        # Build system prompt
        system_msg = (
            "You are an assistant answering questions using the provided document snippets. "
            "Answer concisely and refer only to the snippets. If not present, say you don't know.\n\nSnippets:\n"
            + context
        )
        
        # Load previous messages for memory
        base = os.path.join(SESSIONS_DIR, req.session_id)
        previous_messages = []
        if os.path.isdir(base):
            meta_path = os.path.join(base, 'meta.json')
            try:
                with open(meta_path, 'r', encoding='utf8') as f:
                    meta = json.load(f)
                previous_messages = meta.get('messages', [])
            except Exception:
                previous_messages = []
        
        # Build conversation history (last 5 exchanges for context window efficiency)
        conversation_history = []
        for msg in previous_messages[-10:]:  # Keep last 10 messages (5 exchanges)
            conversation_history.append({'role': msg.get('role', 'user'), 'content': msg.get('text', '')})
        
        # Query Gemini or Ollama with memory
        answer = None
        if GENAI_AVAILABLE:
            try:
                client = genai.Client(api_key=API_KEY)
                # Build full message content with history
                full_content = system_msg + "\n\n"
                if conversation_history:
                    full_content += "Previous conversation:\n"
                    for msg in conversation_history:
                        role = "You" if msg['role'] == 'assistant' else "User"
                        full_content += f"{role}: {msg['content']}\n"
                    full_content += "\n"
                full_content += f"User: {req.question}"
                resp = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=full_content
                )
                answer = resp.text
            except Exception as e:
                print(f"Gemini failed: {e}. Falling back to Ollama...")
                answer = None
        
        # Fallback to Ollama if Gemini failed or unavailable
        if answer is None:
            if OLLAMA_AVAILABLE:
                # Build messages list for Ollama with memory
                messages = [{'role': 'system', 'content': system_msg}]
                messages.extend(conversation_history)
                messages.append({'role': 'user', 'content': req.question})
                
                resp = ollama.chat(
                    model='gemma3:12b',
                    messages=messages
                )
                answer = resp.get('message', {}).get('content', '')
            else:
                raise HTTPException(status_code=500, detail='No LLM available (Gemini and Ollama both failed)')
        
        # Save QA to session meta
        if os.path.isdir(base):
            meta_path = os.path.join(base, 'meta.json')
            with open(meta_path, 'r', encoding='utf8') as f:
                meta = json.load(f)
            msgs = meta.get('messages', [])
            msgs.append({'role':'user', 'text': req.question, 'ts': int(__import__('time').time())})
            msgs.append({'role':'assistant', 'text': answer, 'ts': int(__import__('time').time())})
            meta['messages'] = msgs
            with open(meta_path, 'w', encoding='utf8') as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)
        
        return JSONResponse({'answer': answer, 'retrieved_chunks': retrieved, 'context': context, 'retrieved_count': len(retrieved)})
    
    except HTTPException:
        raise
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + '\n' + tb)


class PDFRequest(BaseModel):
    text: str


@app.post("/generate_pdf")
async def generate_pdf(req: PDFRequest):
    if not REPORTLAB_AVAILABLE:
        raise HTTPException(status_code=500, detail="reportlab not installed. Install with: python -m pip install reportlab")
    if not req.text:
        raise HTTPException(status_code=400, detail="text is required to generate PDF")

    try:
        pdf_bytes = text_to_pdf_bytes(req.text)
        return StreamingResponse(BytesIO(pdf_bytes), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=notes.pdf"})
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + "\n" + tb)


@app.get('/sessions')
async def list_sessions():
    ids = []
    for name in os.listdir(SESSIONS_DIR):
        if os.path.isdir(os.path.join(SESSIONS_DIR,name)):
            ids.append(name)
    return JSONResponse({'sessions': ids})


@app.get('/session/{session_id}')
async def get_session(session_id: str):
    base = os.path.join(SESSIONS_DIR, session_id)
    if not os.path.isdir(base):
        raise HTTPException(status_code=404, detail='session not found')
    with open(os.path.join(base,'meta.json'),'r',encoding='utf8') as f:
        meta = json.load(f)
    with open(os.path.join(base,'chunks.json'),'r',encoding='utf8') as f:
        chunks = json.load(f)
    return JSONResponse({'meta': meta, 'chunks_count': len(chunks)})


class ExportRequest(BaseModel):
    session_id: str
    format: str = 'csv'  # csv | xlsx | docx


@app.post('/export')
async def export_endpoint(req: ExportRequest):
    base = os.path.join(SESSIONS_DIR, req.session_id)
    if not os.path.isdir(base):
        raise HTTPException(status_code=404, detail='session not found')
    with open(os.path.join(base,'meta.json'),'r',encoding='utf8') as f:
        meta = json.load(f)
    with open(os.path.join(base,'chunks.json'),'r',encoding='utf8') as f:
        chunks = json.load(f)
    text = '\n\n'.join(chunks)
    structured = meta.get('structured', {})

    if req.format == 'csv':
        df = pd.DataFrame([structured])
        out = BytesIO(); df.to_csv(out, index=False, encoding='utf-8'); out.seek(0)
        return StreamingResponse(out, media_type='text/csv', headers={'Content-Disposition':'attachment; filename=export.csv'})
    elif req.format == 'xlsx':
        df = pd.DataFrame([structured])
        out = BytesIO(); df.to_excel(out, index=False, engine='openpyxl'); out.seek(0)
        return StreamingResponse(out, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition':'attachment; filename=export.xlsx'})
    elif req.format == 'docx':
        doc = Document(); doc.add_heading('Extracted structured data', level=1)
        for k,v in structured.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_page_break(); doc.add_heading('Full text', level=1); doc.add_paragraph(text)
        out = BytesIO(); doc.save(out); out.seek(0)
        return StreamingResponse(out, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition':'attachment; filename=export.docx'})
    else:
        raise HTTPException(status_code=400, detail='unsupported format')


@app.get('/visualize/{session_id}')
async def visualize(session_id: str, top_k: int = 10):
    base = os.path.join(SESSIONS_DIR, session_id)
    if not os.path.isdir(base):
        raise HTTPException(status_code=404, detail='session not found')
    with open(os.path.join(base,'chunks.json'),'r',encoding='utf8') as f:
        chunks = json.load(f)
    full = ' '.join(chunks).lower()
    # basic tokenization and stopword removal
    stop = set(['the','and','a','to','of','in','for','is','on','that','this','with','as','an','are','it','by'])
    words = [w.strip('.,()[]\"\'"') for w in full.split()]
    freq = {}
    for w in words:
        if not w or w in stop or len(w) < 3: continue
        freq[w] = freq.get(w,0) + 1
    items = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_k]
    return JSONResponse({'top_terms': items})


# If run directly, start uvicorn
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
